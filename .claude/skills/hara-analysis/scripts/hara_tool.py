#!/usr/bin/env python3
"""Utility commands for the Claude HARA analysis skill."""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable


SKILL_DIR = Path(__file__).resolve().parents[1]
ASSET_DIR = SKILL_DIR / "assets"

DERIVE_COLUMNS = [
    "No.",
    "子功能",
    "功能丧失",
    "过大",
    "过早",
    "过小",
    "过晚",
    "非预期激活",
    "卡滞",
    "方向错误",
]

MF_HAZARD_COLUMNS = [
    "No.",
    "Milf_ID",
    "故障描述",
    "整车级危害",
    "备注",
]

HARA_COLUMNS = [
    "List_No",
    "MF_ID",
    "故障描述",
    "整车危害",
    "道路类型",
    "道路条件",
    "环境条件",
    "车辆状态",
    "车速(km/h)",
    "特殊要素",
    "附加条件",
    "驾驶员是否在车上",
    "危害事件",
    "E-解释",
    "暴露频率'E'",
    "有风险的人员",
    "可能的后果('S'的理由)",
    "Severity 'S'",
    "C-解释",
    "控制能力 'C'",
    "结果ASIL",
    "安全目标",
    "安全状态",
    "FTTI(ms)",
]

SG_COLUMNS = [
    "SG_No",
    "安全目标",
    "ASIL Level",
    "安全状态",
    "操作模式",
    "FTTI(ms)",
]

SHEET_COLUMNS = {
    "derive_mf": DERIVE_COLUMNS,
    "mf_vehicle_hazards": MF_HAZARD_COLUMNS,
    "HARA": HARA_COLUMNS,
    "sg_sum": SG_COLUMNS,
}

GUIDEWORD_COLUMNS = DERIVE_COLUMNS[2:]
E_LEVELS = {"E0", "E1", "E2", "E3", "E4"}
S_LEVELS = {"S0", "S1", "S2", "S3"}
C_LEVELS = {"C0", "C1", "C2", "C3"}
ASIL_LEVELS = {"QM", "A", "B", "C", "D"}
ASIL_ORDER = {"QM": 0, "A": 1, "B": 2, "C": 3, "D": 4}
MODE_LEVELS = {"驻车", "行车", "驻车和行车"}

FAST_ROADS = {"高速公路", "高速公路-上/下匝道", "城市快速路"}
URBAN_INTERSECTION_SPECIALS = {"红绿灯路口", "环岛路口"}
VRU_TERMS = {"行人", "横穿马路", "自行车", "骑行者"}
VRU_ALLOWED_CONTEXT_TERMS = {"施工区域", "事故场景", "服务区", "收费站", "应急车道", "路肩", "隧道养护", "故障处置", "救援"}
MOVING_STATES = {
    "前向加速",
    "前向匀速",
    "前向减速/刹车",
    "前向滑行",
    "后向行驶",
    "转弯(左转/右转)",
    "变道",
    "变道-紧急避让",
    "超车",
    "调头",
    "自动巡航",
    "智能领航",
}
PARKING_STATES = {"驻车", "停车熄火", "停车未熄火"}
SLOW_STATES = {"启动起步", "前向爬行(缓慢前进)", "泊车", "自动泊车"}
BACKWARD_MOTION_TERMS = {
    "向后滑", "向后溜", "向后移", "向后移动", "后退", "后溜", "撞击后方", "撞上后方",
    "与后车", "后方行人", "后方车辆", "后方道路使用者", "后方障碍物", "侧后方行人"
}
FORWARD_MOTION_TERMS = {
    "向前滑", "向前溜", "向前移", "向前移动", "前移", "前溜", "撞击前方", "撞上前方",
    "与前方车辆", "与前车", "前方行人", "前方车辆", "前车", "前方障碍物", "进入主路"
}
REVERSE_ORIENTATION_TERMS = {"车头朝上坡", "车头朝坡上", "倒车方向", "倒车", "外力向后", "被前方车辆推动"}
FORWARD_FORCE_TERMS = {"驱动扭矩", "D挡", "前进挡", "蠕行", "外力向前", "被后方车辆", "后方车辆追尾", "被推向前"}


@dataclass
class Finding:
    severity: str
    check: str
    message: str
    row: str = ""


def read_json(path: Path) -> Any:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def write_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="\n") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
        f.write("\n")


def text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def load_allowed_values() -> tuple[dict[str, list[str]], set[str]]:
    scenarios = read_json(ASSET_DIR / "operation_scenarios.json")
    hazards_raw = read_json(ASSET_DIR / "vehicle_hazards.json")
    hazards = {v for values in hazards_raw.values() for v in values}
    return scenarios, hazards


def parse_numeric_suffix(value: str, prefix: str) -> int:
    value = text(value).upper()
    if not re.fullmatch(prefix + r"\d", value):
        raise ValueError(f"invalid level: {value}")
    return int(value[1:])


def calculate_asil(s_level: str, e_level: str, c_level: str) -> str:
    s = parse_numeric_suffix(s_level, "S")
    e = parse_numeric_suffix(e_level, "E")
    c = parse_numeric_suffix(c_level, "C")
    if s == 0 or e == 0 or c == 0:
        return "QM"
    score = s + e + c
    if score <= 6:
        return "QM"
    if score == 7:
        return "A"
    if score == 8:
        return "B"
    if score == 9:
        return "C"
    return "D"


def split_fault_cell(cell: str) -> list[str]:
    """Split a derive_mf cell into normalized MF descriptions."""
    cell = text(cell)
    if not cell:
        return []
    matches = list(re.finditer(r"MF\d{3,}", cell))
    if len(matches) > 1:
        faults = []
        for idx, match in enumerate(matches):
            start = match.start()
            end = matches[idx + 1].start() if idx + 1 < len(matches) else len(cell)
            fault = cell[start:end].strip(" ；;，,。")
            if fault:
                faults.append(fault)
        return faults
    match = re.match(r"^(MF\d{3,})\s*(.+)$", cell)
    if not match:
        return [cell]
    mf_code, body = match.groups()
    body = body.strip()
    # Prefer explicit punctuation. Do not split on Chinese "或": many valid
    # descriptions use it inside one condition, for example "低照度或夜间".
    parts = re.split(r"\s*[；;]\s*", body)
    normalized = []
    for part in parts:
        part = part.strip(" ，,。")
        if part:
            normalized.append(f"{mf_code} {part}")
    return normalized or [cell]


def collect_derive_faults(rows: list[dict[str, Any]]) -> tuple[set[str], set[str]]:
    fault_codes: set[str] = set()
    fault_descriptions: set[str] = set()
    for row in rows:
        for col in GUIDEWORD_COLUMNS:
            for fault in split_fault_cell(text(row.get(col))):
                match = re.match(r"^(MF\d{3,})\b", fault)
                if match:
                    fault_codes.add(match.group(1))
                fault_descriptions.add(fault)
    return fault_codes, fault_descriptions


def flatten_analysis(analysis: dict[str, Any]) -> dict[str, list[dict[str, Any]]]:
    return {
        "derive_mf": list(analysis.get("derive_mf") or []),
        "mf_vehicle_hazards": list(analysis.get("mf_vehicle_hazards") or []),
        "HARA": list(analysis.get("HARA") or []),
        "sg_sum": list(analysis.get("sg_sum") or []),
    }


def add(finding_list: list[Finding], severity: str, check: str, message: str, row: str = "") -> None:
    finding_list.append(Finding(severity=severity, check=check, message=message, row=row))


def validate_columns(
    sheets: dict[str, list[dict[str, Any]]], findings: list[Finding]
) -> None:
    for sheet_name, columns in SHEET_COLUMNS.items():
        rows = sheets.get(sheet_name)
        if rows is None:
            add(findings, "blocking", "sheet", f"缺少 sheet 数据：{sheet_name}")
            continue
        for idx, row in enumerate(rows, 1):
            missing = [col for col in columns if col not in row]
            extra = [col for col in row.keys() if col not in columns]
            if missing:
                add(findings, "blocking", "columns", f"缺少字段：{', '.join(missing)}", f"{sheet_name}[{idx}]")
            if extra:
                add(findings, "warning", "columns", f"存在非输出契约字段：{', '.join(extra)}", f"{sheet_name}[{idx}]")


def validate_ids(system_name: str, sheets: dict[str, list[dict[str, Any]]], findings: list[Finding]) -> None:
    system_prefix = re.escape(system_name) if system_name else r"[A-Za-z0-9_\-\u4e00-\u9fff]+"
    seen: dict[str, set[str]] = defaultdict(set)
    patterns = {
        "derive_mf": ("No.", re.compile(rf"^{system_prefix}_fc\d{{2,}}$")),
        "mf_vehicle_hazards": ("Milf_ID", re.compile(rf"^{system_prefix}_Milf_\d{{3,}}$")),
        "HARA": ("MF_ID", re.compile(rf"^{system_prefix}_MF_\d{{2,}}$")),
        "sg_sum": ("SG_No", re.compile(r"^SG\d{3,}$")),
    }
    for sheet_name, (field, pattern) in patterns.items():
        for idx, row in enumerate(sheets.get(sheet_name, []), 1):
            value = text(row.get(field))
            if not value:
                add(findings, "blocking", "id", f"{field} 为空", f"{sheet_name}[{idx}]")
                continue
            if not pattern.match(value):
                add(findings, "blocking", "id", f"{field} 编号不符合规则：{value}", f"{sheet_name}[{idx}]")
            if value in seen[sheet_name]:
                add(findings, "blocking", "id", f"{field} 编号重复：{value}", f"{sheet_name}[{idx}]")
            seen[sheet_name].add(value)

    for idx, row in enumerate(sheets.get("derive_mf", []), 1):
        non_empty_index = 0
        for col in GUIDEWORD_COLUMNS:
            value = text(row.get(col))
            if not value:
                continue
            non_empty_index += 1
            expected = f"MF{idx}{non_empty_index:02d}"
            if not value.startswith(expected):
                add(
                    findings,
                    "warning",
                    "mf-numbering",
                    f"{col} 建议以 {expected} 开头，当前为：{value[:40]}",
                    f"derive_mf[{idx}]",
                )


def validate_source_functions(
    analysis: dict[str, Any], sheets: dict[str, list[dict[str, Any]]], findings: list[Finding]
) -> None:
    project = analysis.get("project") or {}
    source_functions = [text(x) for x in (project.get("source_functions") or []) if text(x)]
    if not source_functions:
        add(
            findings,
            "warning",
            "source-functions",
            "project.source_functions 为空，无法校验 sheet1 是否严格来自源文档功能清单",
        )
        return
    derive_functions = [text(row.get("子功能")) for row in sheets.get("derive_mf", [])]
    source_set = set(source_functions)
    derive_set = set(derive_functions)
    for value in derive_functions:
        if value and value not in source_set:
            add(
                findings,
                "blocking",
                "source-functions",
                f"sheet1 子功能不在源文档功能清单中：{value}",
                "derive_mf",
            )
    missing = [name for name in source_functions if name not in derive_set]
    if missing:
        add(
            findings,
            "blocking",
            "source-functions",
            f"源文档功能清单中的功能未进入 sheet1：{', '.join(missing)}",
            "derive_mf",
        )


def fault_body(value: str) -> str:
    return re.sub(r"^MF\d{3,}\s*", "", text(value)).strip()


def action_signature(value: str) -> str:
    body = fault_body(value)
    actions = [
        "拉起",
        "释放",
        "点亮",
        "熄灭",
        "打开",
        "关闭",
        "锁止",
        "解锁",
        "加速",
        "减速",
        "制动",
        "转向",
        "升高",
        "降低",
        "伸出",
        "收回",
    ]
    for action in actions:
        if action in body:
            return action
    return ""


def starts_with_condition(value: str) -> bool:
    body = fault_body(value)
    condition_prefixes = [
        "未请求",
        "无请求",
        "不满足",
        "行驶中",
        "静止",
        "低速",
        "高速",
        "驻车时",
        "车速",
        "电源",
        "休眠",
        "网络",
        "整车",
        "应",
        "请求",
        "在",
    ]
    return any(body.startswith(prefix) for prefix in condition_prefixes)


def is_bidirectional_current_function(function_name: str, row_text: str, left: str, right: str) -> bool:
    if left not in row_text or right not in row_text:
        return False
    has_left = left in function_name
    has_right = right in function_name
    return has_left == has_right


def validate_derive_mf_semantics(sheets: dict[str, list[dict[str, Any]]], findings: list[Finding]) -> None:
    """Heuristic sheet1 review focused on the malfunction thinking model."""
    discrete_request_terms = [
        "开关",
        "按钮",
        "请求",
        "拉起",
        "释放",
        "打开",
        "关闭",
        "点亮",
        "熄灭",
        "主动",
        "手动",
    ]
    timing_terms = ["自动", "定时", "时序", "阶段", "周期", "延时", "超时", "保持", "条件满足前", "提前"]
    precondition_terms = [
        "电源",
        "休眠",
        "网络",
        "车速",
        "车辆静止",
        "Released",
        "Applied",
        "ON",
        "OFF",
        "档",
        "条件",
    ]
    effect_dimension_hints = [
        (["灯", "照明", "光"], ["亮度", "照度", "光强"]),
        (["制动", "刹车", "驻车", "EPB", "拉起"], ["制动力", "夹紧力", "保持力"]),
        (["转向", "方向盘", "舵"], ["转角", "转向助力", "横摆响应"]),
        (["驱动", "动力", "加速", "扭矩"], ["扭矩", "驱动力", "加速度"]),
        (["车窗", "天窗", "座椅", "尾门"], ["位置", "速度", "作用力"]),
        (["加热", "冷却", "空调"], ["温度", "风量", "持续时间"]),
    ]
    action_pairs = [
        ("打开", "关闭"),
        ("点亮", "熄灭"),
        ("拉起", "释放"),
        ("锁止", "解锁"),
        ("加速", "减速"),
        ("左", "右"),
        ("前进", "后退"),
        ("升高", "降低"),
        ("伸出", "收回"),
    ]

    for idx, row in enumerate(sheets.get("derive_mf", []), 1):
        function_name = text(row.get("子功能"))
        row_text = " ".join(text(row.get(col)) for col in GUIDEWORD_COLUMNS)
        function_context = f"{function_name} {row_text}"
        all_faults = [fault for col in GUIDEWORD_COLUMNS for fault in split_fault_cell(text(row.get(col)))]
        effects = [
            effect
            for triggers, dimensions in effect_dimension_hints
            if contains_any(function_context, triggers)
            for effect in dimensions
        ]
        present_pairs = [(a, b) for a, b in action_pairs if a in function_context and b in function_context]
        bidirectional_pairs = [
            (a, b) for a, b in present_pairs if is_bidirectional_current_function(function_name, row_text, a, b)
        ]

        loss_faults = split_fault_cell(text(row.get("功能丧失")))
        loss_text = text(row.get("功能丧失"))
        stuck_text = text(row.get("卡滞"))
        if len(loss_faults) > 1 and contains_any(" ".join(loss_faults), precondition_terms):
            add(
                findings,
                "warning",
                "sheet1-malfunction",
                "功能丧失疑似按前置条件重复列出多个冗余故障；sheet1 应概括功能级偏离，前置条件放到 sheet2/HARA",
                f"derive_mf[{idx}]",
            )
        if (
            loss_text
            and stuck_text
            and contains_any(loss_text, ["无法熄灭", "无法关闭", "无法释放", "无法解锁", "无法停止"])
            and not bidirectional_pairs
        ):
            add(
                findings,
                "warning",
                "sheet1-malfunction",
                "功能丧失疑似与卡滞重复；若本质是保持在某状态无法改变，应优先归入卡滞",
                f"derive_mf[{idx}]",
            )
        for left, right in bidirectional_pairs:
            has_left_loss = f"无法{left}" in loss_text
            has_right_loss = f"无法{right}" in loss_text
            if has_left_loss != has_right_loss:
                add(
                    findings,
                    "warning",
                    "sheet1-malfunction",
                    f"双向状态控制功能的功能丧失应覆盖两个 required control actions：{left} 和 {right}",
                    f"derive_mf[{idx}]",
                )

        too_early = text(row.get("过早"))
        if too_early and contains_any(function_context, discrete_request_terms) and not contains_any(too_early, timing_terms):
            add(
                findings,
                "warning",
                "sheet1-malfunction",
                "主动离散请求类功能通常没有“过早”；请求前发生应归入非预期激活",
                f"derive_mf[{idx}]",
            )

        if effects:
            effect_text = "、".join(dict.fromkeys(effects))
            if not text(row.get("过大")):
                add(
                    findings,
                    "warning",
                    "sheet1-malfunction",
                    f"行为模型提示可能存在作用量维度（{effect_text}）；若该作用量在正常功能中存在，应确认是否遗漏“过大”故障",
                    f"derive_mf[{idx}]",
                )
            if not text(row.get("过小")):
                add(
                    findings,
                    "warning",
                    "sheet1-malfunction",
                    f"行为模型提示可能存在作用量维度（{effect_text}）；若该作用量在正常功能中存在，应确认是否遗漏“过小”故障",
                    f"derive_mf[{idx}]",
                )

        if present_pairs and not text(row.get("方向错误")):
            pair_text = "、".join(f"{a}/{b}" for a, b in present_pairs)
            add(
                findings,
                "warning",
                "sheet1-malfunction",
                f"行为模型包含相反状态或方向（{pair_text}）；应确认是否存在“请求 A 却执行 B”的方向错误",
                f"derive_mf[{idx}]",
            )

        unexpected = text(row.get("非预期激活"))
        if unexpected and present_pairs and not any(a in unexpected or b in unexpected for a, b in present_pairs):
            pair_text = "、".join(f"{a}/{b}" for a, b in present_pairs)
            add(
                findings,
                "warning",
                "sheet1-malfunction",
                f"非预期激活描述偏泛；对于相反状态（{pair_text}）建议说明非预期发生的是哪一方向",
                f"derive_mf[{idx}]",
            )
        unexpected_faults = split_fault_cell(unexpected)
        unexpected_actions: dict[str, list[str]] = defaultdict(list)
        for fault in unexpected_faults:
            signature = action_signature(fault)
            if signature:
                unexpected_actions[signature].append(fault)
        duplicated_actions = [action for action, faults in unexpected_actions.items() if len(faults) > 1]
        if duplicated_actions:
            add(
                findings,
                "warning",
                "sheet1-malfunction",
                "非预期激活疑似按操作域或允许条件重复拆分同一本质故障；sheet1 保留一个功能级非预期动作，条件放到 sheet2/HARA",
                f"derive_mf[{idx}]",
            )

        if all_faults and any(starts_with_condition(fault) for fault in all_faults):
            add(
                findings,
                "warning",
                "sheet1-style",
                "sheet1 故障描述疑似混入场景/条件句；建议统一为“功能对象 + 偏离表现”的短语风格",
                f"derive_mf[{idx}]",
            )
        for col in GUIDEWORD_COLUMNS:
            value = text(row.get(col))
            if any(f"{left}/{right}" in value for left, right in action_pairs):
                add(
                    findings,
                    "warning",
                    "sheet1-style",
                    f"{col} 中不同动作/状态不应使用“/”压缩；应使用分号拆成独立故障描述",
                    f"derive_mf[{idx}]",
                )
        stuck = text(row.get("卡滞"))
        state_words = [word for pair in present_pairs for word in pair] + ["Applied", "Released", "开", "关", "启用", "停用"]
        if stuck and not contains_any(stuck, state_words):
            add(findings, "warning", "sheet1-malfunction", "卡滞故障应写明卡滞在具体状态", f"derive_mf[{idx}]")

        direction = text(row.get("方向错误"))
        if direction and not (contains_any(direction, ["请求", "却", "相反", "反向", "错误"]) or any(a in direction and b in direction for a, b in action_pairs)):
            add(findings, "warning", "sheet1-malfunction", "方向错误应写明请求方向与实际执行方向相反", f"derive_mf[{idx}]")


def validate_traceability(
    sheets: dict[str, list[dict[str, Any]]],
    hazards_allowed: set[str],
    findings: list[Finding],
    require_sheet2_completeness: bool = True,
) -> None:
    derive_codes, derive_faults = collect_derive_faults(sheets.get("derive_mf", []))
    sheet2_pairs: set[tuple[str, str]] = set()
    sheet2_faults: set[str] = set()
    for idx, row in enumerate(sheets.get("mf_vehicle_hazards", []), 1):
        fault = text(row.get("故障描述"))
        hazard = text(row.get("整车级危害"))
        code_match = re.match(r"^(MF\d{3,})\b", fault)
        if not code_match:
            add(findings, "blocking", "traceability", f"故障描述未以 MF 编号开头：{fault}", f"mf_vehicle_hazards[{idx}]")
        elif code_match.group(1) not in derive_codes:
            add(
                findings,
                "blocking",
                "traceability",
                f"sheet2 故障编号未在 sheet1 中出现：{fault}",
                f"mf_vehicle_hazards[{idx}]",
            )
        if derive_faults and fault not in derive_faults:
            add(
                findings,
                "warning",
                "traceability",
                "sheet2 故障描述没有和 sheet1 完全一致；若为拆分条目，请确认语义一致",
                f"mf_vehicle_hazards[{idx}]",
            )
        if hazard not in hazards_allowed:
            add(findings, "blocking", "enum", f"整车级危害不在允许值中：{hazard}", f"mf_vehicle_hazards[{idx}]")
        sheet2_pairs.add((fault, hazard))
        sheet2_faults.add(fault)

    if require_sheet2_completeness:
        missing_sheet2_faults = sorted(derive_faults - sheet2_faults)
        for fault in missing_sheet2_faults:
            add(
                findings,
                "blocking",
                "traceability",
                f"sheet1 故障未在 sheet2 中逐条出现：{fault}",
                "mf_vehicle_hazards",
            )

    for idx, row in enumerate(sheets.get("HARA", []), 1):
        fault = text(row.get("故障描述"))
        hazard = text(row.get("整车危害"))
        if (fault, hazard) not in sheet2_pairs:
            add(
                findings,
                "blocking",
                "traceability",
                "HARA 行的故障描述/整车危害不能追溯到 sheet2",
                f"HARA[{idx}]",
            )
        if fault not in sheet2_faults:
            add(findings, "blocking", "traceability", f"HARA 故障未在 sheet2 中出现：{fault}", f"HARA[{idx}]")


def validate_sheet2_semantics(sheets: dict[str, list[dict[str, Any]]], findings: list[Finding]) -> None:
    for idx, row in enumerate(sheets.get("mf_vehicle_hazards", []), 1):
        fault = text(row.get("故障描述"))
        hazard = text(row.get("整车级危害"))
        note = text(row.get("备注"))
        if "EPB" in fault and contains_any(fault, ["夹紧力过大", "制动力过大"]) and hazard != "无危害":
            add(
                findings,
                "warning",
                "sheet2-hazard",
                "EPB静态拉起的夹紧力/制动力过大通常只影响寿命，不应偷换为非预期拉起或非预期减速",
                f"mf_vehicle_hazards[{idx}]",
            )
        if "近光灯非预期点亮" in fault and hazard == "无危害":
            add(
                findings,
                "warning",
                "sheet2-hazard",
                "近光灯非预期点亮不能只按白天判断为无危害；夜间/低照度可能影响其他道路使用者判断",
                f"mf_vehicle_hazards[{idx}]",
            )
        if "近光灯请求点亮却熄灭" in fault and contains_any(note, ["非预期熄灭等效", "等效"]):
            add(
                findings,
                "warning",
                "sheet2-hazard",
                "请求点亮却熄灭是方向错误，不能简单等同于行驶中突然非预期熄灭",
                f"mf_vehicle_hazards[{idx}]",
            )
        if "近光灯请求熄灭却点亮" in fault and hazard != "无危害":
            add(
                findings,
                "warning",
                "sheet2-hazard",
                "请求熄灭却点亮通常无车辆级人身伤害风险，不能简单等同于非预期点亮",
                f"mf_vehicle_hazards[{idx}]",
            )


def validate_enums(
    sheets: dict[str, list[dict[str, Any]]], scenarios_allowed: dict[str, list[str]], findings: list[Finding]
) -> None:
    for idx, row in enumerate(sheets.get("HARA", []), 1):
        for field, allowed in scenarios_allowed.items():
            value = text(row.get(field))
            if value not in allowed:
                add(findings, "blocking", "enum", f"{field} 不在允许值中：{value}", f"HARA[{idx}]")
        driver = text(row.get("驾驶员是否在车上"))
        if driver not in {"是", "否", "不涉及"}:
            add(findings, "blocking", "enum", f"驾驶员是否在车上必须为 是/否/不涉及：{driver}", f"HARA[{idx}]")
        e = text(row.get("暴露频率'E'"))
        s = text(row.get("Severity 'S'"))
        c = text(row.get("控制能力 'C'"))
        asil = text(row.get("结果ASIL")).upper()
        if e not in E_LEVELS:
            add(findings, "blocking", "enum", f"E 等级非法：{e}", f"HARA[{idx}]")
        if s not in S_LEVELS:
            add(findings, "blocking", "enum", f"S 等级非法：{s}", f"HARA[{idx}]")
        if c not in C_LEVELS:
            add(findings, "blocking", "enum", f"C 等级非法：{c}", f"HARA[{idx}]")
        if asil not in ASIL_LEVELS:
            add(findings, "blocking", "enum", f"ASIL 非法：{asil}", f"HARA[{idx}]")

    for idx, row in enumerate(sheets.get("sg_sum", []), 1):
        asil = text(row.get("ASIL Level"))
        mode = text(row.get("操作模式"))
        if asil not in {"A", "B", "C", "D"}:
            add(findings, "blocking", "enum", f"sg_sum 的 ASIL Level 必须为 A/B/C/D：{asil}", f"sg_sum[{idx}]")
        if mode not in MODE_LEVELS:
            add(findings, "blocking", "enum", f"操作模式必须为 驻车/行车/驻车和行车：{mode}", f"sg_sum[{idx}]")


def operation_mode_from_row(row: dict[str, Any]) -> str:
    state = text(row.get("车辆状态"))
    return "驻车" if any(k in state for k in ["驻车", "停车", "泊车"]) else "行车"


def sg_signature(row: dict[str, Any]) -> tuple[str, str, str, str, str]:
    return (
        text(row.get("安全目标")),
        text(row.get("ASIL Level")),
        text(row.get("安全状态")),
        text(row.get("操作模式")),
        text(row.get("FTTI(ms)")),
    )


def validate_asil_and_sg(sheets: dict[str, list[dict[str, Any]]], findings: list[Finding]) -> None:
    rows_by_fault: dict[str, list[tuple[int, dict[str, Any]]]] = defaultdict(list)
    for idx, row in enumerate(sheets.get("HARA", []), 1):
        s = text(row.get("Severity 'S'"))
        e = text(row.get("暴露频率'E'"))
        c = text(row.get("控制能力 'C'"))
        asil = text(row.get("结果ASIL")).upper()
        if s in S_LEVELS and e in E_LEVELS and c in C_LEVELS:
            expected = calculate_asil(s, e, c)
            if asil != expected:
                add(findings, "blocking", "asil", f"ASIL 与 S/E/C 不一致，应为 {expected}，当前为 {asil}", f"HARA[{idx}]")
        safety_goal = text(row.get("安全目标"))
        safe_state = text(row.get("安全状态"))
        ftti = text(row.get("FTTI(ms)"))
        if asil == "QM":
            if safety_goal or safe_state or ftti:
                add(findings, "blocking", "sg", "QM 条目安全目标/安全状态/FTTI 必须为空", f"HARA[{idx}]")
        elif asil in {"A", "B", "C", "D"}:
            if not safety_goal or not safe_state or not ftti:
                add(findings, "blocking", "sg", "非 QM 条目必须填写安全目标、安全状态和 FTTI", f"HARA[{idx}]")
        rows_by_fault[text(row.get("故障描述"))].append((idx, row))

    sg_rows = sheets.get("sg_sum", [])
    for idx, row in enumerate(sg_rows, 1):
        goal = text(row.get("安全目标"))
        if not goal:
            add(findings, "blocking", "sg", "sg_sum 安全目标为空", f"sg_sum[{idx}]")
    expected = Counter()
    for fault, indexed_rows in rows_by_fault.items():
        candidate_idx = None
        candidate_row: dict[str, Any] | None = None
        max_asil = "QM"
        for idx, row in indexed_rows:
            asil = text(row.get("结果ASIL")).upper()
            if asil in ASIL_ORDER and ASIL_ORDER[asil] > ASIL_ORDER[max_asil]:
                max_asil = asil
                candidate_idx = idx
                candidate_row = row
        if max_asil == "QM":
            continue
        assert candidate_row is not None and candidate_idx is not None
        goal = text(candidate_row.get("安全目标"))
        safe_state = text(candidate_row.get("安全状态"))
        ftti = text(candidate_row.get("FTTI(ms)"))
        if not goal or not safe_state or not ftti:
            add(findings, "blocking", "sg", f"故障 {fault} 的最高 ASIL 条目缺少安全目标/安全状态/FTTI", f"HARA[{candidate_idx}]")
            continue
        expected_signature = (goal, max_asil, safe_state, operation_mode_from_row(candidate_row), ftti)
        expected[expected_signature] += 1
        for idx, row in indexed_rows:
            asil = text(row.get("结果ASIL")).upper()
            if asil == "QM":
                continue
            row_signature = (
                text(row.get("安全目标")),
                max_asil,
                text(row.get("安全状态")),
                operation_mode_from_row(candidate_row),
                text(row.get("FTTI(ms)")),
            )
            if row_signature != expected_signature:
                add(findings, "blocking", "sg", f"同一故障 MF 的非 QM 条目应继承最高 ASIL 条目的安全目标/安全状态/FTTI：{fault}", f"HARA[{idx}]")
    actual = Counter(sg_signature(row) for row in sg_rows)
    if actual != expected:
        add(
            findings,
            "blocking",
            "sg",
            "sg_sum 必须按每个故障 MF 的最高 ASIL 条目逐条生成；最高为 QM 的故障不生成 SG，且不得按相同安全目标跨 MF 合并",
        )


def contains_any(value: str, needles: Iterable[str]) -> bool:
    return any(n in value for n in needles)


def speed_upper(speed: str) -> int | None:
    speed = text(speed)
    if speed == "静止状态":
        return 0
    match = re.search(r"(\d+)\s*]", speed.replace(",", "，"))
    if not match:
        return None
    return int(match.group(1))


def scenario_condition_text(row: dict[str, Any]) -> str:
    fields = [
        "道路类型",
        "道路条件",
        "环境条件",
        "车辆状态",
        "车速(km/h)",
        "特殊要素",
        "附加条件",
        "危害事件",
        "有风险的人员",
        "可能的后果('S'的理由)",
        "C-解释",
    ]
    return " ".join(text(row.get(field)) for field in fields)


def scenario_fact_text(row: dict[str, Any]) -> str:
    fields = [
        "道路类型",
        "道路条件",
        "环境条件",
        "车辆状态",
        "车速(km/h)",
        "特殊要素",
        "附加条件",
        "危害事件",
        "有风险的人员",
    ]
    return " ".join(text(row.get(field)) for field in fields)


def sec_reason_text(row: dict[str, Any]) -> str:
    fields = ["E-解释", "可能的后果('S'的理由)", "C-解释"]
    return " ".join(text(row.get(field)) for field in fields)


def has_low_light_context(row: dict[str, Any], facts: str) -> bool:
    environment = text(row.get("环境条件"))
    special = text(row.get("特殊要素"))
    return (
        environment in {"夜间", "雾霾", "降雨(小/大/暴)", "降雪(小/大/暴)"}
        or special == "隧道"
        or contains_any(facts, ["夜间", "傍晚", "黑暗", "无路灯", "路灯灭", "低照度", "隧道", "能见度低"])
    )


def validate_sec_reason_scene_alignment(row: dict[str, Any], idx: int, findings: list[Finding]) -> None:
    facts = scenario_fact_text(row)
    reasons = sec_reason_text(row)
    if not reasons:
        return

    required_terms = [
        ("起步", ["起步", "停车", "驻车", "坡道", "蠕行", "D挡", "临时停车"]),
        ("傍晚", ["傍晚"]),
        ("刚接手车辆", ["刚接手车辆", "接手车辆"]),
        ("误以为已开灯", ["误以为已开灯", "以为已开灯"]),
        ("未及时发现", ["未及时发现", "未能及时发现", "难以及时发现"]),
        ("隧道", ["隧道"]),
        ("无路灯", ["无路灯", "路灯灭", "路灯条件(灭)"]),
        ("施工", ["施工"]),
        ("障碍物", ["障碍物", "遗弃的货物", "遗落"]),
        ("横穿", ["横穿"]),
        ("前车距离较近", ["前车距离较近", "与前车距离较近", "跟车距离较近"]),
        ("跟车距离较近", ["前车距离较近", "与前车距离较近", "跟车距离较近"]),
    ]
    for term, support_terms in required_terms:
        if term in reasons and not contains_any(facts, support_terms):
            add(
                findings,
                "blocking",
                "sec-scene-alignment",
                f"SEC 理由引入了场景事实中没有的前提：{term}",
                f"HARA[{idx}]",
            )

    if contains_any(reasons, ["进入暗环境", "进入低照度", "低照度区域"]) and not (
        has_low_light_context(row, facts) and contains_any(facts, ["进入", "驶入", "切换", "傍晚", "隧道"])
    ):
        add(
            findings,
            "blocking",
            "sec-scene-alignment",
            "SEC 理由使用进入暗环境/低照度区域作为暴露前提，但场景未体现进入、驶入、切换、傍晚或隧道等转换条件",
            f"HARA[{idx}]",
        )

    if "日间" in reasons and text(row.get("环境条件")) != "日间" and "日间" not in facts:
        add(findings, "blocking", "sec-scene-alignment", "SEC 理由引用日间条件，但场景不是日间", f"HARA[{idx}]")
    if "夜间" in reasons and text(row.get("环境条件")) != "夜间" and "夜间" not in facts:
        add(findings, "blocking", "sec-scene-alignment", "SEC 理由引用夜间条件，但场景不是夜间", f"HARA[{idx}]")
    high_speed_as_contrast = contains_any(reasons, ["不会引发高速", "不等同", "非高速"]) or re.search(r"不[^。；;]{0,30}高速", reasons)
    if "高速" in reasons and not high_speed_as_contrast and not contains_any(facts, ["高速公路", "城市快速路", "(60，90]", "(90，120]", "高速"]):
        add(findings, "warning", "sec-scene-alignment", "SEC 理由引用高速条件，但场景道路/车速未体现高速", f"HARA[{idx}]")


def validate_scenario_condition_logic(row: dict[str, Any], idx: int, findings: list[Finding]) -> None:
    fault = text(row.get("故障描述"))
    hazard = text(row.get("整车危害"))
    road = text(row.get("道路类型"))
    road_condition = text(row.get("道路条件"))
    environment = text(row.get("环境条件"))
    vehicle_state = text(row.get("车辆状态"))
    speed = text(row.get("车速(km/h)"))
    special = text(row.get("特殊要素"))
    event = text(row.get("危害事件"))
    risk_people = text(row.get("有风险的人员"))
    severity = text(row.get("Severity 'S'"))
    combined = scenario_condition_text(row)
    facts = scenario_fact_text(row)
    upper_speed = speed_upper(speed)
    is_s0 = severity == "S0"

    # 1) 运动学一致性：车辆状态、车速和驾驶员位置必须能同时成立。
    if vehicle_state in PARKING_STATES and speed not in {"静止状态", "(0，10]", "不涉及"}:
        add(findings, "blocking", "scenario-logic", f"停车/驻车状态 {vehicle_state} 与车速 {speed} 不一致", f"HARA[{idx}]")
    if vehicle_state in MOVING_STATES and speed == "静止状态":
        add(findings, "blocking", "scenario-logic", f"行驶状态 {vehicle_state} 与静止车速不一致", f"HARA[{idx}]")
    if vehicle_state in {"泊车", "自动泊车", "前向爬行(缓慢前进)"} and upper_speed is not None and upper_speed > 10:
        add(findings, "blocking", "scenario-logic", f"{vehicle_state} 通常不应使用 {speed} 车速", f"HARA[{idx}]")
    if vehicle_state == "启动起步" and upper_speed is not None and upper_speed > 30:
        add(findings, "warning", "scenario-logic", f"启动起步场景车速 {speed} 偏高，需确认是否应为前向加速/匀速", f"HARA[{idx}]")
    if vehicle_state in {"变道", "变道-紧急避让", "超车", "自动巡航", "智能领航"} and upper_speed is not None and upper_speed <= 10:
        add(findings, "warning", "scenario-logic", f"{vehicle_state} 与低速/静止车速 {speed} 组合可疑", f"HARA[{idx}]")
    if text(row.get("驾驶员是否在车上")) == "否" and contains_any(combined, ["驾驶员踩", "驾驶员重新操作", "驾驶员立即制动", "驾驶员可通过脚刹"]):
        add(findings, "blocking", "scenario-logic", "驾驶员不在车上，但场景/C解释依赖驾驶员直接踩刹或重新操作", f"HARA[{idx}]")

    # 2) 道路拓扑一致性：道路类型、道路条件、特殊要素和风险对象必须属于同一种道路世界。
    high_speed_road = road in FAST_ROADS
    if high_speed_road and contains_any(combined, VRU_TERMS) and not contains_any(combined, VRU_ALLOWED_CONTEXT_TERMS):
        add(
            findings,
            "blocking",
            "scenario-logic",
            "高速/城市快速路通常为封闭或快速通行道路，不能默认使用行人横穿/骑行者作为危险对象；需改为前车、障碍物，或补充施工/事故/应急车道等前提",
            f"HARA[{idx}]",
        )
    if high_speed_road and special in URBAN_INTERSECTION_SPECIALS and road != "高速公路-上/下匝道":
        add(findings, "blocking", "scenario-logic", f"{road} 与特殊要素 {special} 不匹配", f"HARA[{idx}]")
    if road_condition in {"单行道", "多车道隔离"} and contains_any(combined, ["对向来车", "对向车辆", "迎面", "正面碰撞", "会车"]):
        add(findings, "blocking", "scenario-logic", f"{road_condition} 与对向来车/正面碰撞机制冲突", f"HARA[{idx}]")
    if road_condition == "单行道" and contains_any(combined, ["后向来车", "对向车道"]):
        add(findings, "warning", "scenario-logic", "单行道场景出现对向/后向来车表述，需复核交通流方向", f"HARA[{idx}]")

    # 3) 环境一致性：文本里的天气/光照条件要和环境字段相互支撑。
    tunnel_context = special == "隧道" or "隧道" in facts
    if contains_any(facts, ["夜间", "黑暗", "无路灯", "路灯灭", "低照度"]) and environment == "日间" and not tunnel_context:
        add(findings, "warning", "scenario-logic", "场景文字为夜间/黑暗/低照度，但环境条件为日间", f"HARA[{idx}]")
    if contains_any(facts, ["白天", "日间", "环境光充足", "能见度良好"]) and environment == "夜间":
        add(findings, "warning", "scenario-logic", "场景文字为白天/日间/环境光充足，但环境条件为夜间", f"HARA[{idx}]")
    if contains_any(facts, ["暴雨", "大雨", "降雨", "雨天", "雨后", "积水", "湿滑", "浮水"]) and environment != "降雨(小/大/暴)" and road_condition != "浮水路面":
        add(findings, "warning", "scenario-logic", "场景文字包含降雨/湿滑/积水条件，但环境/道路条件未体现降雨或浮水", f"HARA[{idx}]")
    if contains_any(facts, ["暴雪", "大雪", "降雪", "雪天", "冰雪"]) and environment != "降雪(小/大/暴)" and road_condition != "冰雪路面":
        add(findings, "warning", "scenario-logic", "场景文字包含降雪/冰雪条件，但环境/道路条件未体现降雪或冰雪路面", f"HARA[{idx}]")
    if "隧道" in facts and special not in {"隧道", "ALL"}:
        add(findings, "warning", "scenario-logic", "场景文字包含隧道条件，但特殊要素未标为隧道", f"HARA[{idx}]")

    # 4) 风险对象路径一致性：受害对象必须能位于危害事件路径内。
    if special == "行人数量(无)" and contains_any(event + risk_people, ["行人", "横穿马路"]):
        add(findings, "blocking", "scenario-logic", "特殊要素为行人数量(无)，但危害事件或风险人员包含行人", f"HARA[{idx}]")
    if contains_any(event + risk_people, ["对向车辆", "对向来车"]) and road_condition in {"单行道", "多车道隔离"}:
        add(findings, "blocking", "scenario-logic", f"{road_condition} 场景不应默认存在对向车辆风险对象", f"HARA[{idx}]")
    if contains_any(event, ["追尾", "撞击前车", "前车"]) and not contains_any(facts, ["前车", "前方车辆", "慢行车辆", "停驶车辆", "与前车距离", "后车", "后方车辆", "后方排队车辆", "跟车距离", "交通堵塞", "临近拥堵"]):
        add(findings, "warning", "scenario-logic", "追尾/前车危险事件缺少前方车辆或拥堵等支撑条件", f"HARA[{idx}]")
    if contains_any(event, ["追尾后方", "追尾后车", "追尾后方来车"]):
        add(findings, "blocking", "scenario-logic", "追尾表述的相对方位错误：向后溜移应写撞击后方车辆，不应写追尾后方车辆", f"HARA[{idx}]")

    # 5) 故障-危害-事件机理一致性：危险事件应体现当前整车危害，而不是换成另一种危害。
    if "视野" in hazard and not is_s0 and not contains_any(event + combined, ["视野", "照明", "看清", "发现", "识别", "能见度", "黑暗", "低照度"]):
        add(findings, "warning", "scenario-logic", "整车危害为视野丢失/降低，但危险事件未体现照明或识别能力下降", f"HARA[{idx}]")
    if "误判" in hazard and not is_s0 and not contains_any(event + combined, ["误判", "判断", "眩目", "干扰", "位置", "距离", "意图"]):
        add(findings, "warning", "scenario-logic", "整车危害为其他道路使用者误判，但危险事件未体现误判/眩目/位置距离判断机制", f"HARA[{idx}]")
    if hazard == "非预期的纵向移动" and not is_s0:
        if road_condition not in {"上坡道路", "下坡道路"} and not contains_any(combined, ["坡", "蠕行", "外力", "驱动扭矩", "释放已有保持力"]):
            add(findings, "warning", "scenario-logic", "非预期纵向移动缺少坡度、蠕行、外力、驱动扭矩或已有保持力释放作为力来源", f"HARA[{idx}]")
        if road_condition == "下坡道路" and contains_any(facts, BACKWARD_MOTION_TERMS) and not contains_any(facts, REVERSE_ORIENTATION_TERMS):
            add(
                findings,
                "blocking",
                "scenario-logic",
                "下坡道路失去驻车保持时默认应沿车辆前进方向向前溜移；当前场景写成向后/后方风险对象，除非说明车头朝上坡、倒车方向或反向外力",
                f"HARA[{idx}]",
            )
        if road_condition == "上坡道路" and contains_any(facts, FORWARD_MOTION_TERMS):
            add(
                findings,
                "blocking",
                "scenario-logic",
                "上坡道路失去驻车保持时输出表不应写成向前/前方风险对象；若由驱动蠕行或外力向前造成，应改用与运动方向一致的场景条件并在附加条件说明力源",
                f"HARA[{idx}]",
            )
    if hazard == "非预期的减速" and vehicle_state in PARKING_STATES and speed == "静止状态" and not is_s0:
        add(findings, "warning", "scenario-logic", "静止驻车状态下的非预期制动通常不构成非预期减速危险事件", f"HARA[{idx}]")
    if hazard == "驾驶员获取错误的信息" and not contains_any(event + combined, ["误信", "误认为", "未获知", "未提示", "错误提示", "信息"]):
        add(findings, "warning", "scenario-logic", "整车危害为错误信息，但危险事件未体现错误信息如何改变驾驶员行为", f"HARA[{idx}]")


def validate_scenario_coverage(
    sheets: dict[str, list[dict[str, Any]]], findings: list[Finding], min_scenarios: int
) -> None:
    hara_by_pair: dict[tuple[str, str], list[dict[str, Any]]] = defaultdict(list)
    for row in sheets.get("HARA", []):
        hara_by_pair[(text(row.get("故障描述")), text(row.get("整车危害")))].append(row)
    for idx, row in enumerate(sheets.get("mf_vehicle_hazards", []), 1):
        fault = text(row.get("故障描述"))
        hazard = text(row.get("整车级危害"))
        if hazard == "无危害":
            continue
        rows = hara_by_pair.get((fault, hazard), [])
        if len(rows) < min_scenarios:
            add(
                findings,
                "blocking",
                "scenario-coverage",
                f"非无危害故障至少需要 {min_scenarios} 条有效 HARA 场景，当前 {len(rows)} 条：{fault} / {hazard}",
                f"mf_vehicle_hazards[{idx}]",
            )


def heuristic_review(sheets: dict[str, list[dict[str, Any]]], findings: list[Finding]) -> None:
    scenario_keys: dict[tuple[str, ...], int] = {}
    impl_terms = ["传感器", "算法", "报文", "CAN", "驱动", "诊断", "DTC", "软件", "硬件", "冗余", "电路"]
    for idx, row in enumerate(sheets.get("HARA", []), 1):
        fault = text(row.get("故障描述"))
        hazard = text(row.get("整车危害"))
        road = text(row.get("道路类型"))
        road_condition = text(row.get("道路条件"))
        environment = text(row.get("环境条件"))
        vehicle_state = text(row.get("车辆状态"))
        speed = text(row.get("车速(km/h)"))
        special = text(row.get("特殊要素"))
        extra = text(row.get("附加条件"))
        event = text(row.get("危害事件"))
        e_reason = text(row.get("E-解释"))
        s_reason = text(row.get("可能的后果('S'的理由)"))
        c_reason = text(row.get("C-解释"))
        asil = text(row.get("结果ASIL")).upper()
        safety_goal = text(row.get("安全目标"))

        key = (fault, hazard, road, road_condition, environment, vehicle_state, speed, special, event)
        if key in scenario_keys:
            add(findings, "warning", "duplicate", f"与 HARA[{scenario_keys[key]}] 场景高度重复", f"HARA[{idx}]")
        scenario_keys[key] = idx
        validate_scenario_condition_logic(row, idx, findings)
        validate_sec_reason_scene_alignment(row, idx, findings)

        if vehicle_state in {"驻车", "停车熄火", "停车未熄火"} and speed not in {"静止状态", "(0，10]", "不涉及"}:
            add(findings, "blocking", "scenario", f"车辆状态 {vehicle_state} 与车速 {speed} 明显矛盾", f"HARA[{idx}]")
        moving_states = {
            "前向加速",
            "前向匀速",
            "前向减速/刹车",
            "前向滑行",
            "后向行驶",
            "转弯(左转/右转)",
            "变道",
            "变道-紧急避让",
            "超车",
            "调头",
            "自动巡航",
            "智能领航",
        }
        if vehicle_state in moving_states and speed == "静止状态":
            add(findings, "blocking", "scenario", f"车辆状态 {vehicle_state} 与静止车速明显矛盾", f"HARA[{idx}]")
        if vehicle_state in {"驻车", "停车熄火", "停车未熄火"} and road.startswith("高速公路") and not contains_any(extra, ["应急", "服务区", "匝道", "临停"]):
            add(findings, "warning", "scenario", "高速公路驻车场景需要说明应急车道、服务区或临停条件", f"HARA[{idx}]")
        scenario_text = " ".join([road_condition, environment, special, extra, event, e_reason, s_reason, c_reason])
        high_speed_road = road in {"高速公路", "高速公路-上/下匝道", "城市快速路"}
        if high_speed_road and contains_any(
            scenario_text + " " + text(row.get("有风险的人员")),
            ["行人", "横穿马路", "自行车", "骑行者"],
        ) and not contains_any(
            scenario_text,
            ["施工区域", "事故场景", "服务区", "收费站", "应急车道", "路肩", "隧道养护", "故障处置", "救援"],
        ):
            add(
                findings,
                "blocking",
                "scenario",
                "高速/城市快速路通常为封闭或快速通行道路，不能默认使用行人横穿/骑行者作为危险对象；需改为前车、障碍物，或补充施工/事故/应急车道等前提",
                f"HARA[{idx}]",
            )
        if high_speed_road and special in {"红绿灯路口", "环岛路口"} and road != "高速公路-上/下匝道":
            add(findings, "blocking", "scenario", f"{road} 与特殊要素 {special} 不匹配", f"HARA[{idx}]")
        if contains_any(scenario_text, ["暴雨", "大雨", "降雨", "雨天", "雨后", "积水", "湿滑", "浮水"]) and environment not in {"降雨(小/大/暴)", "夜间"}:
            add(findings, "warning", "scenario", "场景文字包含降雨/湿滑/积水条件，但环境条件未体现降雨或夜间雨后组合", f"HARA[{idx}]")
        if contains_any(scenario_text, ["暴雪", "大雪", "降雪", "雪天"]) and environment != "降雪(小/大/暴)":
            add(findings, "warning", "scenario", "场景文字包含降雪条件，但环境条件未体现降雪", f"HARA[{idx}]")
        if "隧道" in scenario_text and special not in {"隧道", "ALL"}:
            add(findings, "warning", "scenario", "场景文字包含隧道条件，但特殊要素未标为隧道", f"HARA[{idx}]")
        if (
            contains_any(fault, ["近光灯", "前照灯", "灯光"])
            and hazard == "驾驶员视野丢失或者降低"
            and text(row.get("Severity 'S'")) != "S0"
            and text(row.get("结果ASIL")).upper() != "QM"
        ):
            if environment not in {"夜间", "雾霾", "降雨(小/大/暴)", "降雪(小/大/暴)"} and special != "隧道" and "隧道" not in extra:
                add(findings, "warning", "scenario", "照明视野降低场景通常需要夜间、隧道或低能见度条件", f"HARA[{idx}]")
        if "近光灯" in fault:
            strong_visibility_loss = contains_any(
                scenario_text,
                ["突然", "无照明", "无路灯", "黑暗", "隧道", "前车距离较近", "慢行车辆", "障碍物", "横穿", "车道变窄", "反应时间", "制动距离不足"],
            )
            if asil in {"C", "D"} and not contains_any(fault, ["非预期熄灭"]) and not strong_visibility_loss:
                add(
                    findings,
                    "warning",
                    "sec",
                    "近光灯非突然熄灭类故障被评为 C/D，需复核驾驶员是否通常可提前发现、减速或停车",
                    f"HARA[{idx}]",
                )
            if contains_any(fault, ["无法熄灭", "非预期点亮", "卡滞在点亮状态", "请求熄灭却点亮"]) and asil not in {"QM", "A"}:
                add(
                    findings,
                    "warning",
                    "sec",
                    "近光灯持续/非预期点亮通常弱于远光灯，不宜评为高ASIL",
                    f"HARA[{idx}]",
                )
            strong_delay_evidence = contains_any(
                scenario_text,
                ["无照明", "无路灯", "黑暗隧道", "进入隧道", "横穿", "突然出现", "慢行车辆", "障碍物", "前车距离较近", "制动距离", "反应时间"],
            )
            if "点亮响应过晚" in fault and asil not in {"QM", "A"} and not strong_delay_evidence:
                add(
                    findings,
                    "warning",
                    "sec",
                    "近光灯点亮响应过晚通常是短暂延迟，除非场景极端且证据充分，否则不宜评为高ASIL",
                    f"HARA[{idx}]",
                )
        if contains_any(fault, ["EPB", "驻车", "拉起"]) and hazard == "非预期的纵向移动":
            if road_condition not in {"上坡道路", "下坡道路"} and not contains_any(scenario_text, ["坡", "蠕行", "外力", "驱动扭矩"]):
                add(findings, "warning", "scenario", "驻车非预期移动需要坡度、蠕行、外力或驱动扭矩等纵向力来源", f"HARA[{idx}]")
            if text(row.get("Severity 'S'")) != "S0" and "坡" in scenario_text and road_condition not in {"上坡道路", "下坡道路"}:
                add(findings, "warning", "scenario", "驻车溜车场景文本包含坡道/坡度，但道路条件未选上坡或下坡", f"HARA[{idx}]")
            if text(row.get("驾驶员是否在车上")) != "否" and (text(row.get("暴露频率'E'")) == "E4" or text(row.get("控制能力 'C'")) == "C3"):
                add(findings, "warning", "sec", "驾驶员在车上的驻车溜车通常仍可通过脚刹/P挡补救，E4或C3需要非常强的场景证据", f"HARA[{idx}]")

        if len(event) < 12:
            add(findings, "warning", "hazardous-event", "危害事件过短，可能缺少风险对象或伤害机制", f"HARA[{idx}]")
        s0_no_hazard_event = text(row.get("Severity 'S'")) == "S0" and contains_any(event, ["不构成", "未造成", "无人员伤害", "无实际安全影响"])
        if not s0_no_hazard_event and not contains_any(event, ["导致", "造成", "引起", "碰撞", "追尾", "撞击", "伤害", "失控", "视野", "误判"]):
            add(findings, "warning", "hazardous-event", "危害事件未明确说明故障导致的危险后果", f"HARA[{idx}]")
        if text(row.get("Severity 'S'")) == "S0" and contains_any(event, ["导致人员伤害风险", "造成人员伤害风险"]):
            add(findings, "warning", "hazardous-event", "S0 场景不应写成存在人员伤害风险", f"HARA[{idx}]")
        if text(row.get("暴露频率'E'")) == "E0" and contains_any(event + s_reason, ["碰撞", "追尾", "撞击", "无法及时", "受伤", "伤害"]):
            add(findings, "warning", "scenario", "E0 表示场景不可信，不应同时保留碰撞/受伤类危险事件；应改为可信低风险场景或删除危险后果", f"HARA[{idx}]")

        if len(e_reason) < 18 or not contains_any(e_reason, ["频", "每", "常", "偶", "小时", "年", "月", "次", "暴露", "场景"]):
            add(findings, "warning", "sec", "E-解释缺少暴露频率或使用场景依据", f"HARA[{idx}]")
        if text(row.get("Severity 'S'")) != "S0" and (
            len(s_reason) < 18 or not contains_any(s_reason, ["伤", "碰撞", "追尾", "撞击", "行人", "乘员", "死亡", "骨折", "致命", "严重"])
        ):
            add(findings, "warning", "sec", "S 理由缺少伤害对象或伤害机制", f"HARA[{idx}]")
        if text(row.get("Severity 'S'")) != "S0" and (
            len(c_reason) < 18 or not contains_any(c_reason, ["反应", "时间", "感知", "控制", "驾驶员", "避让", "空间", "车上"])
        ):
            add(findings, "warning", "sec", "C-解释缺少感知、反应时间、空间或驾驶员状态依据", f"HARA[{idx}]")

        if asil != "QM" and safety_goal and contains_any(safety_goal, impl_terms):
            add(findings, "warning", "safety-goal", "安全目标疑似包含实现方案词汇，应改为车辆级安全意图", f"HARA[{idx}]")
        if asil != "QM" and text(row.get("FTTI(ms)")) and not re.search(r"\d", text(row.get("FTTI(ms)"))):
            add(findings, "warning", "ftti", "FTTI 建议包含可评审的时间数值", f"HARA[{idx}]")


def validate_analysis(
    analysis: dict[str, Any],
    min_scenarios: int = 3,
    require_scenario_coverage: bool = True,
    require_sheet2_completeness: bool = True,
) -> list[Finding]:
    findings: list[Finding] = []
    project = analysis.get("project") or {}
    system_name = text(project.get("system_name"))
    if not system_name:
        add(findings, "blocking", "project", "project.system_name 不能为空")
    if not text(project.get("item_definition")):
        add(findings, "warning", "project", "project.item_definition 为空或过短")

    sheets = flatten_analysis(analysis)
    scenarios_allowed, hazards_allowed = load_allowed_values()
    validate_columns(sheets, findings)
    validate_source_functions(analysis, sheets, findings)
    validate_ids(system_name, sheets, findings)
    validate_derive_mf_semantics(sheets, findings)
    validate_traceability(
        sheets,
        hazards_allowed,
        findings,
        require_sheet2_completeness=require_sheet2_completeness,
    )
    validate_sheet2_semantics(sheets, findings)
    validate_enums(sheets, scenarios_allowed, findings)
    validate_asil_and_sg(sheets, findings)
    if require_scenario_coverage:
        validate_scenario_coverage(sheets, findings, min_scenarios)
    heuristic_review(sheets, findings)
    return findings


def render_report(analysis: dict[str, Any], findings: list[Finding]) -> str:
    sheets = flatten_analysis(analysis)
    blocking = [f for f in findings if f.severity == "blocking"]
    warnings = [f for f in findings if f.severity == "warning"]
    lines = [
        "# HARA 校验与评审报告",
        "",
        "## 摘要",
        "",
        f"- 结论：{'不通过' if blocking else '通过'}",
        f"- Blocking：{len(blocking)}",
        f"- Warning：{len(warnings)}",
        f"- derive_mf 行数：{len(sheets.get('derive_mf', []))}",
        f"- mf_vehicle_hazards 行数：{len(sheets.get('mf_vehicle_hazards', []))}",
        f"- HARA 行数：{len(sheets.get('HARA', []))}",
        f"- sg_sum 行数：{len(sheets.get('sg_sum', []))}",
        "",
    ]
    if blocking:
        lines.extend(["## Blocking Findings", ""])
        for i, finding in enumerate(blocking, 1):
            where = f" ({finding.row})" if finding.row else ""
            lines.append(f"{i}. [{finding.check}]{where} {finding.message}")
        lines.append("")
    if warnings:
        lines.extend(["## Warnings", ""])
        for i, finding in enumerate(warnings, 1):
            where = f" ({finding.row})" if finding.row else ""
            lines.append(f"{i}. [{finding.check}]{where} {finding.message}")
        lines.append("")

    asil_counts: dict[str, int] = defaultdict(int)
    for row in sheets.get("HARA", []):
        asil_counts[text(row.get("结果ASIL")).upper()] += 1
    if asil_counts:
        lines.extend(["## ASIL 分布", ""])
        for level in ["QM", "A", "B", "C", "D"]:
            lines.append(f"- {level}: {asil_counts.get(level, 0)}")
        lines.append("")

    assumptions = analysis.get("project", {}).get("assumptions") or []
    if assumptions:
        lines.extend(["## 需要工程师确认的假设", ""])
        for item in assumptions:
            lines.append(f"- {item}")
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def load_analysis(path: Path) -> dict[str, Any]:
    data = read_json(path)
    if not isinstance(data, dict):
        raise SystemExit(f"analysis must be a JSON object: {path}")
    return data


def apply_function_catalog(analysis: dict[str, Any], catalog_path: str | None) -> None:
    if not catalog_path:
        return
    catalog = read_json(Path(catalog_path))
    functions = [text(item.get("name")) for item in catalog.get("functions", []) if text(item.get("name"))]
    analysis.setdefault("project", {})["source_functions"] = functions


def cmd_init(args: argparse.Namespace) -> int:
    out = Path(args.out)
    out.mkdir(parents=True, exist_ok=True)
    (out / "source").mkdir(exist_ok=True)
    (out / "drafts").mkdir(exist_ok=True)
    system_name = args.system or Path(args.source).stem
    template = {
        "project": {
            "system_name": system_name,
            "source_file": str(args.source),
            "item_definition": "",
            "source_functions": [],
            "assumptions": [],
        },
        "derive_mf": [],
        "mf_vehicle_hazards": [],
        "HARA": [],
        "sg_sum": [],
    }
    analysis_path = out / "analysis.json"
    if not analysis_path.exists() or args.force:
        write_json(analysis_path, template)
    print(f"initialized: {out}")
    print(f"analysis: {analysis_path}")
    return 0


def cmd_extract_docx(args: argparse.Namespace) -> int:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - environment-specific
        raise SystemExit(f"python-docx is required: {exc}") from exc

    source = Path(args.source)
    out = Path(args.out)
    out.mkdir(parents=True, exist_ok=True)
    doc = Document(str(source))
    lines: list[str] = [f"# {source.name}", ""]
    index: list[dict[str, Any]] = []
    functions: list[dict[str, Any]] = []
    current_section: dict[str, Any] | None = None

    def start_section(title: str, line_no: int) -> None:
        nonlocal current_section
        if current_section is not None:
            current_section["end_line"] = line_no - 1
            index.append(current_section)
        current_section = {"title": title, "start_line": line_no, "end_line": line_no}

    for para in doc.paragraphs:
        value = para.text.strip()
        if not value:
            continue
        line_no = len(lines) + 1
        if re.match(r"^\d+(?:\.\d+)*\s+", value) or (len(index) == 0 and current_section is None):
            start_section(value, line_no)
        lines.append(value)
        lines.append("")

    for t_idx, table in enumerate(doc.tables, 1):
        title = f"表格 {t_idx}"
        line_no = len(lines) + 1
        start_section(title, line_no)
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip().replace("\n", " / ") for cell in row.cells])
        if rows:
            headers = rows[0]
            function_col = next((i for i, h in enumerate(headers) if h == "功能"), None)
            equipped_col = next((i for i, h in enumerate(headers) if "搭载" in h), None)
            if function_col is not None:
                for r_idx, row in enumerate(rows[1:], 2):
                    if function_col >= len(row):
                        continue
                    name = text(row[function_col])
                    if not name:
                        continue
                    if equipped_col is not None and equipped_col < len(row):
                        equipped = text(row[equipped_col])
                        if equipped and equipped not in {"√", "是", "Y", "Yes", "YES"}:
                            continue
                    functions.append({"name": name, "source": title, "row": r_idx})
        if rows:
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
        lines.append("")

    if current_section is not None:
        current_section["end_line"] = len(lines)
        index.append(current_section)

    md_path = out / "source.md"
    idx_path = out / "source_index.json"
    fn_path = out / "function_catalog.json"
    md_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8", newline="\n")
    write_json(idx_path, {"source": str(source), "sections": index})
    write_json(fn_path, {"source": str(source), "functions": functions})
    print(f"extracted markdown: {md_path}")
    print(f"source index: {idx_path}")
    print(f"function catalog: {fn_path}")
    return 0


def cmd_validate(args: argparse.Namespace) -> int:
    analysis = load_analysis(Path(args.analysis))
    apply_function_catalog(analysis, args.function_catalog)
    findings = validate_analysis(
        analysis,
        min_scenarios=args.min_scenarios,
        require_scenario_coverage=not args.skip_scenario_coverage,
        require_sheet2_completeness=not args.skip_sheet2_completeness,
    )
    report = render_report(analysis, findings)
    if args.out:
        Path(args.out).parent.mkdir(parents=True, exist_ok=True)
        Path(args.out).write_text(report, encoding="utf-8", newline="\n")
    print(report)
    return 1 if any(f.severity == "blocking" for f in findings) else 0


def cmd_export(args: argparse.Namespace) -> int:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter
    except Exception as exc:  # pragma: no cover - environment-specific
        raise SystemExit(f"openpyxl is required: {exc}") from exc

    analysis = load_analysis(Path(args.analysis))
    apply_function_catalog(analysis, args.function_catalog)
    findings = validate_analysis(
        analysis,
        min_scenarios=args.min_scenarios,
        require_scenario_coverage=not args.skip_scenario_coverage,
        require_sheet2_completeness=not args.skip_sheet2_completeness,
    )
    report = render_report(analysis, findings)
    if args.review_out:
        Path(args.review_out).parent.mkdir(parents=True, exist_ok=True)
        Path(args.review_out).write_text(report, encoding="utf-8", newline="\n")
    if any(f.severity == "blocking" for f in findings) and not args.allow_blocking:
        print(report)
        print("export blocked: fix blocking findings or pass --allow-blocking", file=sys.stderr)
        return 1

    sheets = flatten_analysis(analysis)
    wb = Workbook()
    default = wb.active
    wb.remove(default)
    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    for sheet_name, columns in SHEET_COLUMNS.items():
        ws = wb.create_sheet(sheet_name)
        ws.append(columns)
        for row in sheets.get(sheet_name, []):
            ws.append([row.get(col, "") for col in columns])
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        for col_idx, col_name in enumerate(columns, 1):
            letter = get_column_letter(col_idx)
            max_len = len(col_name)
            for cell in ws[letter]:
                max_len = max(max_len, min(len(text(cell.value)), 80))
                cell.alignment = Alignment(vertical="top", wrap_text=True)
            ws.column_dimensions[letter].width = min(max(max_len + 2, 10), 48)
        for row_idx in range(2, ws.max_row + 1):
            ws.row_dimensions[row_idx].height = 42

    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out)
    print(f"exported: {out}")
    if args.review_out:
        print(f"review: {args.review_out}")
    return 0


def workbook_to_analysis(path: Path) -> dict[str, Any]:
    try:
        from openpyxl import load_workbook
    except Exception as exc:  # pragma: no cover - environment-specific
        raise SystemExit(f"openpyxl is required: {exc}") from exc
    wb = load_workbook(path, data_only=True)
    analysis = {
        "project": {
            "system_name": path.stem.split("_")[0],
            "source_file": "",
            "item_definition": "从 Excel 反读用于校验",
            "source_functions": [],
            "assumptions": [],
        },
        "derive_mf": [],
        "mf_vehicle_hazards": [],
        "HARA": [],
        "sg_sum": [],
    }
    for sheet_name, columns in SHEET_COLUMNS.items():
        if sheet_name not in wb.sheetnames:
            continue
        ws = wb[sheet_name]
        headers = [text(c.value) for c in ws[1]]
        rows = []
        for raw_row in ws.iter_rows(min_row=2, values_only=True):
            if all(text(v) == "" for v in raw_row):
                continue
            row = {}
            for col in columns:
                if col in headers:
                    row[col] = text(raw_row[headers.index(col)])
                else:
                    row[col] = ""
            rows.append(row)
        analysis[sheet_name] = rows
    if analysis["derive_mf"]:
        first_no = text(analysis["derive_mf"][0].get("No."))
        if "_fc" in first_no:
            analysis["project"]["system_name"] = first_no.split("_fc", 1)[0]
    return analysis


def cmd_review_excel(args: argparse.Namespace) -> int:
    analysis = workbook_to_analysis(Path(args.source))
    apply_function_catalog(analysis, args.function_catalog)
    findings = validate_analysis(
        analysis,
        min_scenarios=args.min_scenarios,
        require_scenario_coverage=not args.skip_scenario_coverage,
        require_sheet2_completeness=not args.skip_sheet2_completeness,
    )
    report = render_report(analysis, findings)
    if args.out:
        Path(args.out).parent.mkdir(parents=True, exist_ok=True)
        Path(args.out).write_text(report, encoding="utf-8", newline="\n")
    print(report)
    return 1 if any(f.severity == "blocking" for f in findings) else 0


def cmd_asil(args: argparse.Namespace) -> int:
    print(calculate_asil(args.s, args.e, args.c))
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="HARA skill deterministic utilities")
    sub = parser.add_subparsers(dest="command", required=True)

    p_init = sub.add_parser("init", help="create a HARA working directory")
    p_init.add_argument("--source", required=True)
    p_init.add_argument("--out", required=True)
    p_init.add_argument("--system")
    p_init.add_argument("--force", action="store_true")
    p_init.set_defaults(func=cmd_init)

    p_extract = sub.add_parser("extract-docx", help="extract DOCX text and tables to markdown")
    p_extract.add_argument("--source", required=True)
    p_extract.add_argument("--out", required=True)
    p_extract.set_defaults(func=cmd_extract_docx)

    p_validate = sub.add_parser("validate", help="validate analysis JSON")
    p_validate.add_argument("--analysis", required=True)
    p_validate.add_argument("--out")
    p_validate.add_argument("--function-catalog")
    p_validate.add_argument("--strict", action="store_true", help="reserved for prompt readability; blocking findings fail")
    p_validate.add_argument("--min-scenarios", type=int, default=3)
    p_validate.add_argument("--skip-scenario-coverage", action="store_true")
    p_validate.add_argument("--skip-sheet2-completeness", action="store_true")
    p_validate.set_defaults(func=cmd_validate)

    p_export = sub.add_parser("export", help="export analysis JSON to Excel")
    p_export.add_argument("--analysis", required=True)
    p_export.add_argument("--out", required=True)
    p_export.add_argument("--review-out")
    p_export.add_argument("--function-catalog")
    p_export.add_argument("--allow-blocking", action="store_true")
    p_export.add_argument("--min-scenarios", type=int, default=3)
    p_export.add_argument("--skip-scenario-coverage", action="store_true")
    p_export.add_argument("--skip-sheet2-completeness", action="store_true")
    p_export.set_defaults(func=cmd_export)

    p_review = sub.add_parser("review-excel", help="review an exported HARA workbook")
    p_review.add_argument("--source", required=True)
    p_review.add_argument("--out")
    p_review.add_argument("--function-catalog")
    p_review.add_argument("--min-scenarios", type=int, default=3)
    p_review.add_argument("--skip-scenario-coverage", action="store_true")
    p_review.add_argument("--skip-sheet2-completeness", action="store_true")
    p_review.set_defaults(func=cmd_review_excel)

    p_asil = sub.add_parser("asil", help="calculate ASIL from S/E/C")
    p_asil.add_argument("--s", required=True)
    p_asil.add_argument("--e", required=True)
    p_asil.add_argument("--c", required=True)
    p_asil.set_defaults(func=cmd_asil)

    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    return args.func(args)


if __name__ == "__main__":
    raise SystemExit(main())
